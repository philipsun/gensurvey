#-------------------------------------------------------------------------------
# Name:        模块1
# Purpose:
#
# Author:      sunxp
#
# Created:     11/10/2020
# Copyright:   (c) sunxp 2020
# Licence:     <your licence>
#-------------------------------------------------------------------------------
import re
from docx import Document
def makedocx(fname,sentlist):
    document = Document()
    document.add_heading('Advances in Text Summarization', 0)  #插入标题
    p = document.add_paragraph('An automatic generated summarization from a set of papers ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    docsent = ""
    addpara = True
    for eachsent in sentlist:
        if re.search('the authors present following contributions',eachsent):
            if docsent:
                document.add_paragraph(docsent + '. ')
                docsent = ""
            docsent = eachsent
        elif re.search('\d-t',eachsent):
            document.add_heading(eachsent,1)
        elif re.search('REFERENCE',eachsent):
           document.add_heading(eachsent,1)
           addpara = False
        elif addpara:
            docsent = docsent +eachsent
        else:
            if docsent:
                document.add_paragraph(docsent)
                docsent = ""
            document.add_paragraph(eachsent + '. ')

    if docsent:
        document.add_paragraph(docsent + '. ')
    document.save(fname+'.docx')
def main():
    fn = r'E:\pythonworknew\code\textsum\test\survey\gen_survey_wordquery_allv6ks_dual_sentrank_opt_num.txt'
    file = open(fn, 'r',encoding='utf-8')
    try:
        text_lines = file.readlines()
      #  print(type(text_lines), text_lines)
        for line in text_lines:
            print(type(line), line)

        makedocx(fn,text_lines)
    finally:
        file.close()

if __name__ == '__main__':
    main()
